"""Uvoz papirnatih nesukladnosti (Word obrasci) u QMS.

Izvor: mapa u kojoj svaka nesukladnost ima svoju pod-mapu, a u njoj Word obrazac
„Nesukladnost" (tablica od ~33 reda) plus slike, mailovi i ostali prilozi.

Pokretanje (iz mape appa, jer se privici spremaju u .\\uploads):

    .venv\\Scripts\\python.exe uvoz_nesukladnosti.py "C:\\...\\2026 Nesukladnosti"
    .venv\\Scripts\\python.exe uvoz_nesukladnosti.py "C:\\...\\2026 Nesukladnosti" --proba

`--proba` sve ispiše, ali ništa ne upisuje.

Ponovno pokretanje je sigurno: nesukladnost koja već postoji (isti broj predmeta)
se preskače, pa se uvoz može dopuniti kad stignu nove.
"""
from __future__ import annotations

import re
import sys
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from sqlalchemy import select

from app.core import migracije
from app.core.database import SessionLocal
from app.modules.reklamacije import models as _m  # noqa: F401 — registrira tablice
from app.modules.reklamacije import privici as priv
from app.modules.reklamacije.models import CAPA, Privitak, Reklamacija, StavkaTroska

migracije.primijeni()   # baza mora imati nove stupce i kad server ne radi

GODINA = 2026
PREFIKS = "NS"                       # NS-01/2026 … razlikuje se od RK- koje app sam dodjeljuje
PRESKOCI = {"thumbs.db", "desktop.ini"}


# ─────────────────────────────────────────────────────────────────────────────
# Čitanje Word obrasca
# ─────────────────────────────────────────────────────────────────────────────

def _oznacena(celija) -> bool:
    """Je li u ćeliji kvačica.

    Obrazac koristi novije Wordove kontrole (`w14:checkbox`), stariji primjerci
    znaju imati staru kontrolu (`w:checkBox`) ili samo znak. Provjeravaju se sve
    tri, jer o tome ovisi hoće li se pročitati Manja/Veća i DA/NE.
    """
    if 'w14:checked w14:val="1"' in celija._tc.xml:
        return True
    for cb in celija._tc.iter(qn("w:checkBox")):
        for ch in cb.iter(qn("w:checked")):
            if ch.get(qn("w:val")) in (None, "1", "true"):
                return True
        for df in cb.iter(qn("w:default")):
            if df.get(qn("w:val")) in ("1", "true"):
                return True
    return any(z in celija.text for z in ("\u2612", "\u2611", "\u25a0", "\u2714"))


def _redci(tablica) -> list[list[str]]:
    """Redci tablice; spojene ćelije se ne ponavljaju, kvačica se označi s [X]."""
    out = []
    for r in tablica.rows:
        celije, zadnji = [], None
        for c in r.cells:
            if c._tc is zadnji:
                continue
            zadnji = c._tc
            celije.append((("[X] " if _oznacena(c) else "") + c.text.strip()).strip())
        out.append(celije)
    return out


def procitaj_obrazac(put: Path) -> dict:
    d = Document(str(put))
    if not d.tables:
        return {}
    redci = _redci(d.tables[0])
    ravno = ["  ||  ".join(r) for r in redci]

    def iza(kljuc: str) -> list[str]:
        for i, r in enumerate(ravno):
            if kljuc.lower() in r.lower():
                return redci[i + 1] if i + 1 < len(redci) else []
        return []

    def uz(kljuc: str) -> str:
        for r in redci:
            for i, c in enumerate(r):
                if kljuc.lower() in c.lower() and i + 1 < len(r):
                    return r[i + 1].strip()
        return ""

    def redak(kljuc: str) -> str:
        return next((x for x in ravno if kljuc.lower() in x.lower()), "")

    zag = iza("Redni broj nesukladnosti")
    proiz = iza("Naziv i oznaka proizvoda")
    odg = iza("Odgovorna osoba za provedbu")

    # Porijeklo: dva reda kvačica ispod naslova, plus slobodan tekst uz „Ostalo".
    i_por = next((i for i, x in enumerate(ravno) if "Porijeklo nesukladnosti" in x), None)
    por_tekst = "\n".join(ravno[i_por + 1: i_por + 3]) if i_por is not None else ""

    return {
        "broj": zag[0] if len(zag) > 0 else "",
        "datum": zag[1] if len(zag) > 1 else "",
        "evidentirao": zag[2] if len(zag) > 2 else "",
        "porijeklo_tekst": por_tekst,
        "proizvod": proiz[0] if len(proiz) > 0 else "",
        "radni_nalog": proiz[1] if len(proiz) > 1 else "",
        "kategorija_red": redak("Kategorija nesukladnosti"),
        "opis": "\n".join(x for x in iza("Opis nesukladnosti") if x),
        "korekcija": "\n".join(x for x in iza("Korekcija") if x),
        "korekciju_proveo": uz("osoba koja je provela korekciju"),
        "uzrok": "\n".join(x for x in iza("Uzrok nesukladnosti") if x),
        "uzrok_potvrdio": uz("odgovorna osoba podru"),
        "vezana": uz("upisati broj iste"),
        "korektivna_red": redak("potrebno provesti korektivnu radnju"),
        "korektivna_radnja": "\n".join(x for x in iza("Potrebna korektivna radnja") if x),
        "odgovorna_osoba": odg[0] if len(odg) > 0 else "",
        "rok": odg[1] if len(odg) > 1 else "",
        "radnju_definirao": uz("koja je definirala korektivnu radnju"),
        "provedena_radnja": "\n".join(x for x in iza("Opis provedene korektivne radnje") if x),
        "radnju_proveo": uz("osoba koja je provela korektivnu radnju"),
        "ucinkovitost": "\n".join(x for x in iza("Provjera u\u010dinkovitosti") if x),
        "ucinkovitost_provjerio": uz("osoba koja je provjerila u"),
        "promjene_red": redak("promjene u sustavu upravljanja"),
        "broj_promjene": uz("Evidencija promjena"),
        "direktor": uz("Direktor"),
        "cijeli_obrazac": "\n".join(ravno),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Pretvorbe
# ─────────────────────────────────────────────────────────────────────────────

def _datum(v: str) -> date | None:
    v = (v or "").strip().rstrip(".")
    for oblik in ("%d.%m.%Y", "%d.%m.%y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(v, oblik).date()
        except ValueError:
            continue
    return None


def _oznaceno_iza(red: str, pojam: str) -> bool:
    """Je li kvačica baš ispred zadanog pojma u tom retku."""
    return bool(re.search(r"\[X\]\s*" + re.escape(pojam), red, re.IGNORECASE))


def _prvo(uzorak: str, tekst: str, skupina: int = 1) -> str | None:
    m = re.search(uzorak, tekst, re.IGNORECASE)
    return m.group(skupina).strip(" .,;") if m else None


def iz_opisa(opis: str) -> dict:
    """Podaci koji su u obrascu utipkani unutar opisa (stroj, materijal, trošak…)."""
    stroj = _prvo(r"Stroj\s*:\s*([^\n]+)", opis)
    strojar = _prvo(r"Strojar\s*:\s*([^\n]+)", opis)
    operater = _prvo(r"Operater\s*:\s*([^\n]+)", opis)
    osoblje = ", ".join(x for x in (strojar, operater) if x) or None
    trosak = _prvo(r"TRO\u0160KOVI\s*:?\s*([\d.]+,\d{2}|\d+[\d.,]*)\s*(?:EUR|\u20ac)", opis)
    return {
        "stroj": stroj,
        "osoblje": osoblje[:200] if osoblje else None,
        "sifra_materijala": _prvo(r"\u0160ifra\s*:?\s*([0-9]{6,})", opis),
        "lot_sarza": _prvo(r"Lot\s*:?\s*([^\n,;]+)", opis),
        "broj_primke": _prvo(r"(?:PRMKA|PRIMKA|primka)\s*:?\s*([^\n,;]+)", opis),
        "dobavljac": _prvo(r"Dobavlja\u010d\s*:?\s*([^\n]+)", opis),
        "trosak": trosak,
    }


def _u_broj(tekst: str) -> float | None:
    """„538,37" -> 538.37 (hrvatski zapis: točka tisućice, zarez decimale)."""
    if not tekst:
        return None
    try:
        return float(tekst.replace(".", "").replace(",", "."))
    except ValueError:
        return None


def naslov_iz_mape(ime_mape: str) -> str:
    """Iz „23-26 RN 967 Omotnica Attiva 40 Omsa" -> „RN 967 Omotnica Attiva 40 Omsa"."""
    t = re.sub(r"^\s*\d+\s*[-/]\s*\d+\s*", "", ime_mape)
    t = re.sub(r"^\s*(Int\.?\s*nes|Int\s*nes|ns)\s*", "", t, flags=re.IGNORECASE)
    t = t.replace("_", " ").strip(" -–")
    return re.sub(r"\s{2,}", " ", t) or ime_mape


def broj_mape(ime_mape: str) -> int | None:
    m = re.match(r"\s*(\d{1,3})\s*[-/]\s*26", ime_mape)
    return int(m.group(1)) if m else None


# ─────────────────────────────────────────────────────────────────────────────
# Uvoz
# ─────────────────────────────────────────────────────────────────────────────

def _skupine(baza: Path) -> dict[int, list[Path]]:
    """Mape grupirane po rednom broju (ista nesukladnost zna imati dvije mape)."""
    out: dict[int, list[Path]] = {}
    for mapa in sorted(p for p in baza.iterdir() if p.is_dir()):
        n = broj_mape(mapa.name)
        if n is None:
            print(f"  ! preskačem mapu bez rednog broja: {mapa.name}")
            continue
        out.setdefault(n, []).append(mapa)
    return dict(sorted(out.items()))


def ime_za_privitak(p: Path) -> str:
    """Ime datoteke za QMS.

    U arhivi ima datoteka nazvanih samo ".msg" (bez imena). Za Path je to ime bez
    nastavka (kao skrivena datoteka), pa bi je provjera tipa odbila. Takvima se
    doda ime, da se sadržaj ne izgubi.
    """
    if p.suffix == "" and p.name.startswith("."):
        return f"poruka{p.name}"
    return p.name


def _privici_za(mape: list[Path]) -> list[Path]:
    dat = []
    for m in mape:
        for p in sorted(m.rglob("*")):
            if p.is_file() and p.name.lower() not in PRESKOCI and not p.name.startswith("~$"):
                dat.append(p)
    return dat


def dopuni_privitke(db, r: Reklamacija, mape: list[Path], upozorenja: list[str]) -> int:
    """Prikvači sve datoteke iz mapa koje uz tu nesukladnost još ne postoje.

    Usporedba ide po imenu i veličini, pa se ponovnim pokretanjem ništa ne duplicira,
    a datoteka koja je prvi put pala (prevelika, čudno ime) se naknadno dohvati.
    """
    vec = {(p.original_name.lower(), p.velicina) for p in r.privici}
    dodano = 0
    for datoteka in _privici_za(mape):
        ime = ime_za_privitak(datoteka)
        try:
            velicina = datoteka.stat().st_size
        except OSError as e:
            upozorenja.append(f"{r.broj_predmeta}: {datoteka.name} — {e}")
            continue
        if (ime.lower(), velicina) in vec:
            continue
        try:
            podaci = datoteka.read_bytes()
        except OSError as e:
            upozorenja.append(f"{r.broj_predmeta}: ne mogu pročitati {datoteka.name} ({e})")
            continue
        rez = priv.spremi_datoteku(r.id, ime, podaci)
        if isinstance(rez, str):
            upozorenja.append(f"{r.broj_predmeta}: {ime} — {rez}")
            continue
        db.add(Privitak(**rez))
        vec.add((ime.lower(), velicina))
        dodano += 1
    return dodano


def uvezi(baza: Path, proba: bool = False) -> None:
    db = SessionLocal()
    novih = privitaka = preskocenih = capa_n = trosak_n = 0
    upozorenja: list[str] = []
    try:
        for n, mape in _skupine(baza).items():
            broj_predmeta = f"{PREFIKS}-{n:02d}/{GODINA}"
            postojeca = db.scalar(
                select(Reklamacija).where(Reklamacija.broj_predmeta == broj_predmeta))
            if postojeca is not None:
                preskocenih += 1
                if not proba:
                    dodano = dopuni_privitke(db, postojeca, mape, upozorenja)
                    if dodano:
                        db.commit()
                        privitaka += dodano
                        print(f"  {broj_predmeta}  već postoji, dodano priloga: {dodano}")
                continue

            docxi = [d for m in mape for d in sorted(m.glob("*.docx"))
                     if not d.name.startswith("~$")]
            o = procitaj_obrazac(docxi[0]) if docxi else {}
            if not docxi:
                upozorenja.append(f"{broj_predmeta}: nema Word obrasca, upisano iz naziva mape "
                                  f"i priloga ({mape[0].name})")

            opis = o.get("opis", "")
            iz = iz_opisa(opis)
            kat_red = o.get("kategorija_red", "")
            kategorija = ("VECA" if _oznaceno_iza(kat_red, "Ve") else
                          "MANJA" if _oznaceno_iza(kat_red, "Manja") else None)
            kor_red = o.get("korektivna_red", "")
            potrebna = ("DA" if _oznaceno_iza(kor_red, "DA") else
                        "NE" if _oznaceno_iza(kor_red, "NE") else None)
            prom_red = o.get("promjene_red", "")
            promjene = ("DA" if _oznaceno_iza(prom_red, "DA") else
                        "NE" if _oznaceno_iza(prom_red, "NE") else None)

            por_tekst = o.get("porijeklo_tekst", "")
            if _oznaceno_iza(por_tekst, "Pogre"):
                porijeklo, por_nap = "POGRESKA_U_RADU", None
            elif _oznaceno_iza(por_tekst, "Ostalo"):
                porijeklo = "OSTALO"
                por_nap = (por_tekst.split("Ostalo", 1)[-1]
                           .split("\n", 1)[-1].strip(" |\n")[:200] or None)
            elif _oznaceno_iza(por_tekst, "Reklamacija"):
                porijeklo, por_nap = "REKLAMACIJA", None
            elif _oznaceno_iza(por_tekst, "Izvje"):
                porijeklo, por_nap = "AUDIT", None
            else:
                porijeklo, por_nap = None, None

            # Materijal kao porijeklo znači da je krivac dobavljač papira.
            izvor = "DOBAVLJAC" if (por_nap and "materijal" in por_nap.lower()) else (
                "INTERNO" if porijeklo == "POGRESKA_U_RADU" else None)

            datum = _datum(o.get("datum", "")) or date(GODINA, 1, 1)
            status = "RIJESENO" if potrebna == "NE" else "U_OBRADI"
            naslov = naslov_iz_mape(mape[0].name)[:200]

            r = Reklamacija(
                broj_predmeta=broj_predmeta,
                vrsta="INTERNA",
                status=status,
                prioritet="VISOK" if kategorija == "VECA" else "SREDNJI",
                kategorija=kategorija,
                izvor=izvor,
                tezina="VELIKI" if kategorija == "VECA" else "MALI",
                naslov=naslov,
                opis=opis or ("(Word obrazac nije popunjen. Podaci su u prilozima "
                              f"iz mape: {mape[0].name})"),
                prijavitelj=(o.get("evidentirao") or "").strip()[:100] or "(nije upisano)",
                kupac_dobavljac=(iz["dobavljac"] or None),
                naziv_proizvoda=(o.get("proizvod") or None),
                broj_radnog_naloga=(o.get("radni_nalog") or None),
                stroj=iz["stroj"],
                osoblje=iz["osoblje"],
                datum_prijave=datetime.combine(datum, datetime.min.time()),
                rok_rjesavanja=_datum(o.get("rok", "")),
                korekcija=(o.get("korekcija") or None),
                analiza_uzroka=(o.get("uzrok") or None),
                vezana_nesukladnost=(o.get("vezana") or None),
                promjene_sustava=promjene,
                broj_promjene=(o.get("broj_promjene") or None),
                ucinkovitost_biljeska=(o.get("ucinkovitost") or None),
                izvorni_broj=(o.get("broj") or None),
                izvorna_mapa=" + ".join(m.name for m in mape)[:255],
                porijeklo=porijeklo,
                porijeklo_napomena=por_nap,
                korekciju_proveo=(o.get("korekciju_proveo") or None),
                uzrok_potvrdio=(o.get("uzrok_potvrdio") or None),
                potrebna_korektivna=potrebna,
                radnju_definirao=(o.get("radnju_definirao") or None),
                sifra_materijala=iz["sifra_materijala"],
                lot_sarza=iz["lot_sarza"],
                broj_primke=iz["broj_primke"],
                direktor_potpis=(o.get("direktor") or None),
                izvorni_zapis=(o.get("cijeli_obrazac") or None),
                napomena=f"Uvezeno iz arhive: {' + '.join(m.name for m in mape)}",
            )

            if proba:
                print(f"  {broj_predmeta}  {naslov[:60]:62} kat={kategorija or '-':5} "
                      f"kor={potrebna or '-':3} priloga={len(_privici_za(mape))}")
                continue

            db.add(r)
            db.flush()
            novih += 1

            # Korektivna mjera iz obrasca
            if (o.get("korektivna_radnja") or o.get("odgovorna_osoba") or o.get("rok")):
                db.add(CAPA(
                    reklamacija_id=r.id,
                    vrsta="KOREKTIVNA",
                    opis_mjere=(o.get("korektivna_radnja")
                                or "(mjera nije opisana u obrascu, upisani su samo nositelj i rok)"),
                    odgovorna_osoba=(o.get("odgovorna_osoba") or "(nije upisano)")[:100],
                    rok_izvrsenja=_datum(o.get("rok", "")),
                    status="IZVRSENA" if o.get("provedena_radnja") else "PLANIRANA",
                    rezultat=(o.get("provedena_radnja") or None),
                    provjerio=(o.get("ucinkovitost_provjerio") or None),
                ))
                capa_n += 1

            # Trošak iz obrasca (COPQ)
            iznos = _u_broj(iz["trosak"] or "")
            if iznos:
                db.add(StavkaTroska(
                    reklamacija_id=r.id, kategorija="OSTALO",
                    opis="Trošak naveden u obrascu nesukladnosti",
                    kolicina=1.0, jedinica="EUR", jed_cijena=iznos, iznos=iznos,
                    tko_snosi="DOBAVLJAC" if izvor == "DOBAVLJAC" else "INTERNO",
                ))
                trosak_n += 1

            # Prilozi: Word obrazac, slike, mailovi, sve iz mape
            privitaka += dopuni_privitke(db, r, mape, upozorenja)

            db.commit()
            print(f"  {broj_predmeta}  {naslov[:58]:60} priloga: {len(_privici_za(mape)):2}")

        if not proba:
            db.commit()
    finally:
        db.close()

    print(f"\nNovih nesukladnosti: {novih}   preskočeno (već postoje): {preskocenih}")
    print(f"Korektivnih mjera: {capa_n}   troškova: {trosak_n}   privitaka: {privitaka}")
    if upozorenja:
        print(f"\nUpozorenja ({len(upozorenja)}):")
        for u in upozorenja:
            print("  -", u)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        raise SystemExit(2)
    izvor = Path(sys.argv[1])
    if not izvor.is_dir():
        print(f"Nema mape: {izvor}")
        raise SystemExit(2)
    uvezi(izvor, proba="--proba" in sys.argv)
